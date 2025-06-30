import streamlit as st
import pandas as pd
import openai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO

# --- Streamlit App Configuration ---
st.set_page_config(page_title="PAR Writing Assistant", layout="wide")
st.title("📋 Feedback Note Generator")
st.markdown("Answer the following questions to generate a structured feedback note.")

# --- Feedback Form ---
with st.form("feedback_form"):
    rank = st.selectbox("Select rank level", ["Cpl", "MCpl", "Sgt", "WO"], index=1)
    last_name = st.text_input("Enter last name of the member")

    event_description = st.text_area("Event Description", placeholder="Brief summary of the event or project")
    q1 = st.text_area("Who was involved and what was done?",
                      placeholder="Describe who completed the task, who was impacted, and what actions were taken.")
    q2 = st.text_area("Where and why did this happen?",
                      placeholder="Explain the location or context and the reason the task or project was necessary.")
    q3 = st.text_area("How was it done and what was the outcome?",
                      placeholder="Highlight any notable processes, tools, or innovations used and the final result.")

    submitted = st.form_submit_button("Generate Feedback Note")

# --- Load Competencies from Excel based on selected rank ---
def load_rank_competencies(selected_rank):
    try:
        df = pd.read_excel("PAR Writing Guide (1).xlsx", sheet_name=selected_rank, engine="openpyxl")
        if 'Competency' in df.columns and 'Facets' in df.columns:
            return df[['Competency', 'Facets']].dropna()
        else:
            st.error(f"The sheet for {selected_rank} does not contain the expected columns.")
            return pd.DataFrame()
    except Exception as e:
        st.error(f"Error loading competencies for {selected_rank}: {e}")
        return pd.DataFrame()

# --- Generate Feedback Note ---
if submitted:
    competency_df = load_rank_competencies(rank)

    if competency_df.empty:
        st.stop()

    # Generate all valid competency:facets pairs
    valid_pairs = [f"{row['Competency'].strip()}: {row['Facet'].strip()}" for _, row in competency_df.iterrows()]
    valid_pairs_text = "; ".join(valid_pairs)

    # Build the prompt for OpenAI
    prompt = f"""
Rank: {rank}
Last Name: {last_name}
Event Description: {event_description}
Who and What: {q1}
Where and Why: {q2}
How and Outcome: {q3}

Using the input above, generate a formal military-style feedback note.

Start with 3–5 valid competencies and facets from the list below. Use only the following format for each:

Competency: Facet (Score) – rationale

Only use valid competency:facet pairs from this list:
{valid_pairs_text}

Do not invent new combinations. Do not use competencies from other ranks. Avoid generic phrasing like 'Leadership' without a facet.

Structure:
Event Description:
[3–5 competency/facet lines formatted as above]
[1–2 paragraph description using the member's rank and last name on first mention, then only they/them/their pronouns thereafter.]

Outcome:
[2–3 sentence measurable or strategic result.]

Use abbreviated rank format (e.g., MCpl Macpherson), and do not spell out ranks. Use they/them pronouns consistently after the first mention.
"""

    try:
        client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a military admin assistant trained to write professional performance feedback notes. "
                        "Always use abbreviated rank formats (e.g., MCpl, WO). Do not spell out full rank names. "
                        "Use the format 'Rank Lastname' on first mention, then use they/them pronouns."
                    )
                },
                {"role": "user", "content": prompt}
            ]
        )

        output_text = response.choices[0].message.content
        st.markdown("### ✏️ Generated Feedback Note")
        st.text_area("Output", output_text, height=400)

        # --- Create Word Document ---
        def create_word_doc(text):
            doc = Document()
            doc.add_heading("Performance Feedback Note", 0)
            lines = text.strip().splitlines()
            for line in lines:
                line = line.strip()
                if line.lower().startswith("event description"):
                    doc.add_paragraph("Event Description:", style='Heading 2')
                elif line.lower().startswith("outcome"):
                    doc.add_paragraph("Outcome:", style='Heading 2')
                elif not line:
                    doc.add_paragraph("")
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(11)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        doc_buffer = create_word_doc(output_text)

        # --- Download Button ---
        st.download_button(
            label="📄 Download as Word Document",
            data=doc_buffer,
            file_name="feedback_note.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"Failed to generate feedback note: {e}")